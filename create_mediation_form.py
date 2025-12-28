from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')

    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{edge}')
        border.set(ns.qn('w:val'), 'single')
        border.set(ns.qn('w:sz'), '8')
        border.set(ns.qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)

# Create document
doc = Document()

# Title
p = doc.add_paragraph("FORM ‘A’")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

p = doc.add_paragraph("MEDIATION APPLICATION FORM\n[REFER RULE 3(1)]")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True

p = doc.add_paragraph("Mumbai District Legal Services Authority\nCity Civil Court, Mumbai")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\nDETAILS OF PARTIES:")

# Create table
table = doc.add_table(rows=12, cols=3)
set_table_borders(table)

data = [
    ("1", "Name of Applicant", "{{client_name}}"),
    ("", "Address and contact details of Applicant",
     "REGISTERED ADDRESS:\n{{branch_address}}\n\nCORRESPONDENCE ADDRESS:\n{{branch_address}}"),
    ("", "Telephone No.", "{{mobile}}"),
    ("", "Mobile No.", "{{mobile}}"),
    ("", "Email ID", "info@kslegal.co.in"),
    ("2", "Opposite Party Details", ""),
    ("", "Name", "{{customer_name}}"),
    ("", "Address",
     "REGISTERED ADDRESS:\n{{address}}\n\nCORRESPONDENCE ADDRESS:\n{{address}}"),
    ("", "Telephone No.", ""),
    ("", "Mobile No.", ""),
    ("", "Email ID", ""),
    ("", "DETAILS OF DISPUTE",
     "Commercial Courts (Pre-Institution Mediation) Rules, 2018"),
]

for i, row in enumerate(data):
    for j, value in enumerate(row):
        cell = table.cell(i, j)
        cell.text = value
        cell.paragraphs[0].runs[0].font.size = Pt(10)

# Save file
doc.save("mediation_application_form.docx")
print("Word document created successfully.")
